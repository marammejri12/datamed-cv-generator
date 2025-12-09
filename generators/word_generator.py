"""
Word Generator - Generate CVs in DOCX format
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


class WordGenerator:
    """Generate anonymous CV in Word format"""

    def __init__(self, template_type: str = 'advanced'):
        self.template_type = template_type

        # Color palettes
        if template_type == 'fastorgie':
            self.primary_color = RGBColor(196, 30, 58)      # Red
            self.dark_color = RGBColor(139, 26, 46)         # Dark red
            self.light_bg = RGBColor(255, 230, 234)         # Light red
        else:  # advanced (DataMed)
            self.primary_color = RGBColor(26, 54, 93)       # Navy blue
            self.dark_color = RGBColor(15, 40, 71)          # Dark navy
            self.light_bg = RGBColor(230, 242, 255)         # Light blue

    def generate_cv(self, data: dict, output_path: str):
        """Generate CV in Word format"""
        doc = Document()

        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.8)
            section.left_margin = Cm(1.8)
            section.right_margin = Cm(1.8)

        # Add logo
        self._add_logo(doc)

        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("Nom & Prénom")
        run.font.size = Pt(28)
        run.font.color.rgb = self.dark_color
        run.bold = True

        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("Consultant IT")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(107, 114, 128)

        doc.add_paragraph()  # Spacing

        # Section counter
        section_num = 1

        # 1. FORMATIONS & CERTIFICATIONS
        if data.get('diplomes'):
            self._add_section_header(doc, section_num, "FORMATIONS & CERTIFICATIONS")
            self._add_diplomes(doc, data['diplomes'])

            if data.get('certifications'):
                self._add_certifications(doc, data['certifications'])

            section_num += 1

        # 2. COMPÉTENCES
        if data.get('competences_groups'):
            self._add_section_header(doc, section_num, "COMPÉTENCES FONCTIONNELLES & TECHNIQUES")
            self._add_competences(doc, data['competences_groups'])
            section_num += 1

        # 3. LANGUES
        if data.get('langues'):
            self._add_section_header(doc, section_num, "LANGUES")
            self._add_langues(doc, data['langues'])
            section_num += 1

        # Page break
        doc.add_page_break()
        self._add_logo(doc)

        # 4. EXPÉRIENCES
        if data.get('experiences'):
            self._add_section_header(doc, section_num, "EXPÉRIENCES PROFESSIONNELLES")
            self._add_experiences(doc, data['experiences'])

        # Footer
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("www.consultingdatamed.com" if self.template_type == 'advanced' else "www.fastor.com")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(107, 114, 128)

        # Save
        doc.save(output_path)
        return output_path

    def _add_logo(self, doc):
        """Add logo"""
        base_path = os.path.dirname(os.path.dirname(__file__))

        if self.template_type == 'fastorgie':
            logo_paths = [
                os.path.join(base_path, 'image', 'fastorgie.png'),
                os.path.join(base_path, 'image', 'fastor.png'),
            ]
        else:
            logo_paths = [
                os.path.join(base_path, 'image', 'datamed_consulting_logo.jfif'),
                os.path.join(base_path, 'image', 'datamed_consulting_logo.png'),
            ]

        for logo_path in logo_paths:
            if os.path.exists(logo_path):
                try:
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    run.add_picture(logo_path, width=Cm(5))
                    doc.add_paragraph()  # Spacing
                    break
                except Exception as e:
                    print(f"Error loading logo: {e}")

    def _add_section_header(self, doc, number: int, title: str):
        """Add section header with number"""
        paragraph = doc.add_paragraph()

        # Add shading to paragraph
        self._set_cell_background(paragraph, self.primary_color)

        # Number
        run = paragraph.add_run(f"{number}  ")
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.bold = True

        # Title
        run = paragraph.add_run(title)
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.bold = True

        doc.add_paragraph()  # Spacing

    def _add_diplomes(self, doc, diplomes: list):
        """Add diplomas"""
        for dip in diplomes:
            annee = dip.get('annee', '')
            diplome = dip.get('diplome', '')
            etablissement = dip.get('etablissement', '')

            p = doc.add_paragraph()

            # Year
            run = p.add_run(f"{annee}    ")
            run.font.size = Pt(12)
            run.font.color.rgb = self.primary_color
            run.bold = True

            # Diploma
            run = p.add_run(f"{diplome}\n")
            run.font.size = Pt(10)
            run.bold = True

            # Institution
            run = p.add_run(f"        {etablissement}")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(107, 114, 128)

    def _add_certifications(self, doc, certifications: list):
        """Add certifications"""
        for cert in certifications:
            annee = cert.get('annee', '')
            nom = cert.get('nom', '')
            organisme = cert.get('organisme', '')

            p = doc.add_paragraph()

            # Year
            run = p.add_run(f"{annee}    ")
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(217, 119, 6)  # Orange
            run.bold = True

            # Certification
            run = p.add_run(f"{nom}\n")
            run.font.size = Pt(10)
            run.bold = True

            # Organisme
            run = p.add_run(f"        {organisme}")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(107, 114, 128)

    def _add_competences(self, doc, competences_groups: list):
        """Add competences"""
        for group in competences_groups:
            categorie = group.get('categorie', '')
            competences = group.get('competences', [])

            if not competences:
                continue

            p = doc.add_paragraph()

            # Category
            run = p.add_run(f"{categorie}:  ")
            run.font.size = Pt(10)
            run.bold = True

            # Skills
            skill_text = ' • '.join(competences)
            run = p.add_run(skill_text)
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(107, 114, 128)

    def _add_langues(self, doc, langues: list):
        """Add languages"""
        for langue in langues:
            nom = langue.get('langue', '')
            niveau = langue.get('niveau', '')

            p = doc.add_paragraph()

            # Language
            run = p.add_run(f"{nom}:  ")
            run.font.size = Pt(10)
            run.bold = True

            # Level
            run = p.add_run(niveau)
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(107, 114, 128)

    def _add_experiences(self, doc, experiences: list):
        """Add experiences"""
        for exp in experiences:
            entreprise = exp.get('entreprise', '')
            periode = exp.get('periode', '')
            poste = exp.get('poste', '')
            lieu = exp.get('lieu', '')

            # Header
            p = doc.add_paragraph()
            self._set_cell_background(p, self.primary_color)

            run = p.add_run(f"{entreprise}    {periode}")
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True

            # Poste
            if poste:
                p = doc.add_paragraph()
                run = p.add_run(poste)
                run.font.size = Pt(11)
                run.bold = True

            # Lieu
            if lieu:
                p = doc.add_paragraph()
                run = p.add_run(f"📍 {lieu}")
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(107, 114, 128)
                run.italic = True

            # Projets
            if exp.get('projets'):
                p = doc.add_paragraph()
                run = p.add_run("Contexte du projet:")
                run.bold = True

                for projet in exp['projets']:
                    if projet.strip():
                        p = doc.add_paragraph(projet, style='List Bullet')
                        p.paragraph_format.left_indent = Inches(0.5)

            # Réalisations
            if exp.get('realisations'):
                p = doc.add_paragraph()
                run = p.add_run("Réalisations:")
                run.bold = True

                for real in exp['realisations']:
                    if real.strip():
                        p = doc.add_paragraph(f"✓ {real}", style='List Bullet')
                        p.paragraph_format.left_indent = Inches(0.5)

            # Environnement
            if exp.get('environnement'):
                env_text = ' • '.join(exp['environnement']) if isinstance(exp['environnement'], list) else str(exp['environnement'])
                p = doc.add_paragraph()
                run = p.add_run("Environnement technique: ")
                run.bold = True
                run = p.add_run(env_text)
                run.font.size = Pt(9.5)

            doc.add_paragraph()  # Spacing

    def _set_cell_background(self, paragraph, color: RGBColor):
        """Set background color for paragraph"""
        shading_elm = OxmlElement('w:shd')
        # RGBColor stores values directly as r, g, b attributes
        hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
        shading_elm.set(qn('w:fill'), hex_color)
        paragraph._element.get_or_add_pPr().append(shading_elm)
